#!/usr/bin/env python3
"""
CML Transcriptie Tool - WhisperX CLI Transcription for macOS Intel
"""

# MUST be first import - patches huggingface/pytorch compatibility
import patch_compat  # noqa: F401

import os
import sys
import subprocess
import tempfile
import shutil
import time
import io
import re
import gc
from pathlib import Path
from datetime import datetime

# --- CONSTANTS ---

SUPPORTED_AUDIO = ['.mp3', '.wav', '.m4a', '.ogg', '.flac', '.aac']
SUPPORTED_VIDEO = ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv']
SUPPORTED_FORMATS = SUPPORTED_AUDIO + SUPPORTED_VIDEO

LANGUAGES = {
    '1':  ('nl', 'Nederlands'),
    '2':  ('en', 'English'),
    '3':  ('fr', 'Francais'),
    '4':  ('de', 'Deutsch'),
    '5':  ('es', 'Espanol'),
    '6':  ('it', 'Italiano'),
    '7':  ('pt', 'Portugues'),
    '8':  ('ja', 'Japanese'),
    '9':  ('zh', 'Chinese'),
    '10': ('ko', 'Korean'),
}

MODELS = {
    '1': ('tiny',     '39M params   - Snelst, basis kwaliteit'),
    '2': ('base',     '74M params   - Snel, redelijke kwaliteit'),
    '3': ('small',    '244M params  - Goede kwaliteit'),
    '4': ('medium',   '769M params  - Zeer goed (aanbevolen)'),
    '5': ('large',    '1550M params - Beste kwaliteit, langzaam'),
    '6': ('large-v3', '1550M params - Nieuwste, beste voor NL'),
}

# Real-time factors for CPU float32 on Intel Mac (slower than Apple Silicon)
RTF_MAP = {
    'tiny': 0.6,
    'base': 1.0,
    'small': 1.6,
    'medium': 3.0,
    'large': 5.0,
    'large-v3': 5.0,
}


# --- PROGRESS DISPLAY ---

class ProgressCapture:
    """
    Captures WhisperX's print_progress output and displays it as a
    formatted progress bar in the terminal.

    WhisperX prints lines like: "Progress: 45.67%..."
    We intercept these and render: [===============>          ] 45.67%
    """

    def __init__(self, phase_name, phase_start_pct, phase_end_pct):
        self.phase_name = phase_name
        self.phase_start = phase_start_pct
        self.phase_end = phase_end_pct
        self.original_stdout = sys.stdout

    def write(self, text):
        match = re.search(r'(\d+\.?\d*)%', text)
        if match:
            local_pct = float(match.group(1))
            overall_pct = self.phase_start + (local_pct / 100.0) * (self.phase_end - self.phase_start)
            self._render_bar(overall_pct)

    def flush(self):
        pass

    def _render_bar(self, pct):
        bar_width = 40
        filled = int(bar_width * pct / 100)
        if filled < bar_width:
            bar = '=' * filled + '>' + ' ' * (bar_width - filled - 1)
        else:
            bar = '=' * bar_width
        self.original_stdout.write(
            f'\r  {self.phase_name}: [{bar}] {pct:6.2f}%'
        )
        self.original_stdout.flush()

    def finish(self):
        self._render_bar(self.phase_end)
        self.original_stdout.write('\n')
        self.original_stdout.flush()


def print_progress_bar(label, pct):
    """Print a standalone progress bar line."""
    bar_width = 40
    filled = int(bar_width * pct / 100)
    if filled < bar_width:
        bar = '=' * filled + '>' + ' ' * (bar_width - filled - 1)
    else:
        bar = '=' * bar_width
    sys.stdout.write(f'\r  {label}: [{bar}] {pct:6.2f}%')
    sys.stdout.flush()


def finish_progress_bar():
    """Move to next line after a progress bar."""
    sys.stdout.write('\n')
    sys.stdout.flush()


# --- UTILITIES ---

def format_timestamp(seconds):
    """Convert seconds to HH:MM:SS or MM:SS format."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def get_duration_ffprobe(filepath):
    """Get media duration in seconds using ffprobe."""
    try:
        result = subprocess.run(
            ['ffprobe', '-v', 'quiet', '-show_entries',
             'format=duration', '-of', 'csv=p=0', filepath],
            capture_output=True, text=True
        )
        return float(result.stdout.strip())
    except Exception:
        return None


def estimate_processing_time(duration_seconds, model_size):
    """Estimate processing time and return formatted string."""
    rtf = RTF_MAP.get(model_size, 3.0)
    estimated = duration_seconds * rtf
    if estimated < 60:
        return f"~{int(estimated)} seconden"
    elif estimated < 3600:
        return f"~{int(estimated / 60)} minuten"
    else:
        return f"~{estimated / 3600:.1f} uur"


# --- FILE DIALOG ---

def select_file_dialog():
    """Open a native macOS file selection dialog using osascript."""
    extensions = [ext.lstrip('.') for ext in SUPPORTED_FORMATS]
    type_list = ', '.join(f'"{ext}"' for ext in extensions)

    script = (
        f'POSIX path of (choose file with prompt '
        f'"Selecteer een audio- of videobestand" '
        f'of type {{{type_list}}})'
    )

    try:
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True, text=True, timeout=300
        )
        if result.returncode == 0:
            path = result.stdout.strip()
            if path and os.path.isfile(path):
                return path
        return None
    except subprocess.TimeoutExpired:
        return None
    except Exception as e:
        print(f"  Fout bij bestandsselectie: {e}")
        return None


def select_file_manual():
    """Fallback: ask user to type/paste a file path."""
    print("  Voer het pad naar het audio/video bestand in:")
    path = input("  Pad: ").strip().strip("'\"")
    if os.path.isfile(path):
        ext = Path(path).suffix.lower()
        if ext in SUPPORTED_FORMATS:
            return path
        else:
            print(f"  Niet-ondersteund formaat: {ext}")
    else:
        print("  Bestand niet gevonden.")
    return None


# --- AUDIO/VIDEO PROCESSING ---

def is_video(filepath):
    """Check if file is a video format."""
    return Path(filepath).suffix.lower() in SUPPORTED_VIDEO


def convert_video_to_audio(video_path, temp_dir):
    """
    Convert video to 16kHz mono WAV using ffmpeg.
    Shows progress via ffmpeg stderr parsing.
    """
    output_path = os.path.join(temp_dir, Path(video_path).stem + '.wav')
    duration = get_duration_ffprobe(video_path)

    cmd = [
        'ffmpeg', '-i', video_path,
        '-vn',
        '-acodec', 'pcm_s16le',
        '-ar', '16000',
        '-ac', '1',
        '-y',
        '-progress', 'pipe:2',
        output_path
    ]

    process = subprocess.Popen(
        cmd, stderr=subprocess.PIPE, stdout=subprocess.DEVNULL,
        universal_newlines=True
    )

    for line in process.stderr:
        time_match = re.search(r'time=(\d+):(\d+):(\d+\.\d+)', line)
        if time_match and duration and duration > 0:
            h, m, s = time_match.groups()
            current = int(h) * 3600 + int(m) * 60 + float(s)
            pct = min(current / duration * 100, 100)
            overall = 5.0 + (pct / 100.0) * 10.0
            print_progress_bar("Video conversie", overall)

    process.wait()
    if process.returncode != 0:
        raise Exception("ffmpeg video conversie mislukt")

    print_progress_bar("Video conversie", 15.00)
    finish_progress_bar()
    return output_path


def compress_audio_if_needed(audio_path, temp_dir):
    """If audio file > 500MB, resample to 16kHz mono WAV."""
    file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)

    if file_size_mb < 500:
        return audio_path

    print(f"  Groot bestand ({file_size_mb:.0f}MB), wordt gecomprimeerd...")
    output_path = os.path.join(temp_dir, 'compressed_audio.wav')

    subprocess.run([
        'ffmpeg', '-i', audio_path,
        '-ar', '16000', '-ac', '1',
        '-acodec', 'pcm_s16le',
        '-y', output_path
    ], capture_output=True, check=True)

    new_size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"  Gecomprimeerd: {file_size_mb:.0f}MB -> {new_size_mb:.0f}MB")
    return output_path


# --- TRANSCRIPTION ---

def run_transcription(audio_path, model_size, language):
    """
    Run the full WhisperX pipeline:
    1. Load model
    2. Transcribe (with VAD + batched inference)
    3. Align (word-level timestamps)

    Returns (segments, detected_language).
    """
    import whisperx
    import torch

    device = "cpu"
    compute_type = "float32"  # Intel CPU does not support int8 well via CTranslate2
    batch_size = 8  # Lower batch size for Intel CPU performance

    # Phase 1: Load model (0-10%)
    print_progress_bar("Model laden", 0.00)

    model = whisperx.load_model(
        model_size,
        device,
        compute_type=compute_type,
        language=language
    )

    print_progress_bar("Model laden", 10.00)
    finish_progress_bar()

    # Phase 2: Load audio (10-15%)
    print_progress_bar("Audio laden", 10.00)
    audio = whisperx.load_audio(audio_path)
    print_progress_bar("Audio laden", 15.00)
    finish_progress_bar()

    # Phase 3: Transcribe with progress (15-70%)
    progress = ProgressCapture("Transcriptie", 15.0, 70.0)
    old_stdout = sys.stdout
    sys.stdout = progress

    try:
        result = model.transcribe(
            audio,
            batch_size=batch_size,
            language=language,
            print_progress=True
        )
    finally:
        sys.stdout = old_stdout

    progress.finish()

    segments = result["segments"]
    detected_language = result.get("language", language)

    print(f"  Taal: {detected_language}, Segmenten: {len(segments)}")

    # Phase 4: Alignment (70-90%)
    align_progress = ProgressCapture("Uitlijning", 70.0, 90.0)
    old_stdout = sys.stdout
    sys.stdout = align_progress

    model_a = None
    try:
        model_a, metadata = whisperx.load_align_model(
            language_code=detected_language,
            device=device
        )
        aligned = whisperx.align(
            segments,
            model_a,
            metadata,
            audio,
            device,
            return_char_alignments=False,
            print_progress=True
        )
    except Exception as e:
        sys.stdout = old_stdout
        print(f"\n  Waarschuwing: Uitlijning mislukt ({e}), ruwe segmenten worden gebruikt.")
        aligned = {"segments": segments}
    else:
        sys.stdout = old_stdout

    align_progress.finish()

    # Cleanup
    del model
    if model_a is not None:
        del model_a
    gc.collect()

    return aligned["segments"], detected_language


# --- WORD EXPORT ---

def export_to_word(segments, metadata, output_path):
    """Export transcript segments to a Word .docx document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_text = metadata.get("title", "Transcriptie")
    title = doc.add_paragraph(title_text)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Metadata table
    meta_heading = doc.add_paragraph('Informatie')
    meta_heading.runs[0].font.size = Pt(14)
    meta_heading.runs[0].font.bold = True
    meta_heading.runs[0].font.name = 'Calibri'

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'

    meta_items = [
        ("Bestand", metadata.get("filename", "")),
        ("Duur", metadata.get("duration_str", "")),
        ("Model", metadata.get("model", "")),
        ("Taal", metadata.get("language", "")),
        ("Datum", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]

    for key, value in meta_items:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = str(value)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(11)

    doc.add_paragraph()

    # Transcript heading
    t_heading = doc.add_paragraph('Transcriptie')
    t_heading.runs[0].font.size = Pt(14)
    t_heading.runs[0].font.bold = True
    t_heading.runs[0].font.name = 'Calibri'

    doc.add_paragraph()

    # Transcript content
    for segment in segments:
        text = segment.get("text", "").strip()
        start = segment.get("start", 0)

        if not text:
            continue

        ts = format_timestamp(start)

        p = doc.add_paragraph()

        ts_run = p.add_run(f"[{ts}] ")
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(128, 128, 128)
        ts_run.font.name = 'Calibri'

        text_run = p.add_run(text)
        text_run.font.size = Pt(11)
        text_run.font.name = 'Calibri'

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Gegenereerd op {datetime.now().strftime('%d/%m/%Y %H:%M')} "
        f"met CML Transcriptie Tool"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.font.name = 'Calibri'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


# --- MAIN ---

def main():
    print()
    print("============================================")
    print("  CML Transcriptie Tool")
    print("  WhisperX Audio/Video Transcriptie")
    print("============================================")
    print()

    # --- Question 1: Language ---
    print("Kies de taal van het audiobestand:")
    print()
    for key, (code, name) in LANGUAGES.items():
        default = " (standaard)" if code == "nl" else ""
        print(f"  {key:>2}. {name} ({code}){default}")
    print()

    lang_choice = input("  Taal [1]: ").strip() or "1"
    if lang_choice not in LANGUAGES:
        print("  Ongeldige keuze, Nederlands wordt gebruikt.")
        lang_choice = "1"

    language_code, language_name = LANGUAGES[lang_choice]
    print(f"  -> {language_name}")
    print()

    # --- Question 2: Model ---
    print("Kies het Whisper model:")
    print()
    for key, (name, desc) in MODELS.items():
        default = " (standaard)" if name == "medium" else ""
        print(f"  {key}. {name:10s} - {desc}{default}")
    print()

    model_choice = input("  Model [4]: ").strip() or "4"
    if model_choice not in MODELS:
        print("  Ongeldige keuze, 'medium' wordt gebruikt.")
        model_choice = "4"

    model_size, model_desc = MODELS[model_choice]
    print(f"  -> {model_size}")
    print()

    # --- File Selection ---
    print("Selecteer een audio- of videobestand...")
    print("(Een Finder venster wordt geopend)")
    print()

    file_path = select_file_dialog()
    if file_path is None:
        print("  Geen bestand geselecteerd via Finder.")
        file_path = select_file_manual()

    if file_path is None:
        print("\n  Geen bestand geselecteerd. Programma wordt afgesloten.")
        return

    file_name = Path(file_path).name
    file_stem = Path(file_path).stem
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
    print(f"  Bestand: {file_name} ({file_size_mb:.1f} MB)")

    # --- File Info & Time Estimate ---
    duration = get_duration_ffprobe(file_path)
    if duration:
        duration_str = format_timestamp(duration)
        est = estimate_processing_time(duration, model_size)
        print(f"  Duur: {duration_str}")
        print(f"  Geschatte verwerkingstijd: {est}")
    else:
        duration_str = "Onbekend"

    print(f"  Device: CPU (Intel)")
    print(f"  Compute: float32")
    print()

    # --- Processing ---
    temp_dir = tempfile.mkdtemp(prefix="cml_transcriptie_")

    try:
        start_time = time.time()

        print("Verwerking gestart...")
        print()

        # Step 1: Video conversion if needed
        audio_path = file_path
        if is_video(file_path):
            print("  Video gedetecteerd, audio extraheren...")
            audio_path = convert_video_to_audio(file_path, temp_dir)

        # Step 2: Compress if needed
        audio_path = compress_audio_if_needed(audio_path, temp_dir)

        # Step 3: Get duration from audio (more accurate after conversion)
        audio_duration = get_duration_ffprobe(audio_path)
        if audio_duration:
            duration_str = format_timestamp(audio_duration)

        # Step 4: Transcribe + Align
        segments, detected_lang = run_transcription(
            audio_path, model_size, language_code
        )

        # Step 5: Export to Word
        print_progress_bar("Word export", 90.00)

        output_dir = Path.home() / "Downloads"
        output_path = output_dir / f"{file_stem}.docx"

        # Handle filename collision
        counter = 1
        while output_path.exists():
            output_path = output_dir / f"{file_stem}_{counter}.docx"
            counter += 1

        metadata = {
            "title": file_stem,
            "filename": file_name,
            "duration_str": duration_str,
            "model": model_size,
            "language": language_name,
        }

        export_to_word(segments, metadata, str(output_path))

        print_progress_bar("Word export", 100.00)
        finish_progress_bar()

        # --- Done ---
        elapsed = time.time() - start_time
        elapsed_str = format_timestamp(elapsed)

        print()
        print("============================================")
        print("  VOLTOOID!")
        print(f"  Verwerkingstijd: {elapsed_str}")
        print(f"  Output: {output_path}")
        print("============================================")
        print()

        # Open Finder to show the file
        subprocess.run(['open', '-R', str(output_path)], capture_output=True)

    except KeyboardInterrupt:
        print("\n\n  Afgebroken door gebruiker.")
    except Exception as e:
        print(f"\n\n  FOUT: {e}")
        import traceback
        traceback.print_exc()
    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            pass


if __name__ == "__main__":
    main()
